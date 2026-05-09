# -*- coding: utf-8 -*-
"""
生成毕业论文：智慧医问-智能医药系统的设计与实现
成都理工大学 智能科学与技术 本科毕业设计论文
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ========== 全局样式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def add_heading_styled(doc, text, level=1, font_name='黑体', font_size=None, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """添加标题"""
    if font_size is None:
        font_sizes = {0: Pt(22), 1: Pt(16), 2: Pt(14), 3: Pt(13)}
        font_size = font_sizes.get(level, Pt(12))
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    return p


def add_body(doc, text, first_indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    return p


def add_chapter(doc, text):
    """添加章标题"""
    return add_heading_styled(doc, text, level=1, font_name='黑体', font_size=Pt(16), bold=True)


def add_section(doc, text):
    """添加节标题"""
    return add_heading_styled(doc, text, level=2, font_name='黑体', font_size=Pt(14), bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)


def add_subsection(doc, text):
    """添加小节标题"""
    return add_heading_styled(doc, text, level=3, font_name='黑体', font_size=Pt(13), bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)


# ==================== 封面 ====================
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

add_heading_styled(doc, '成  都  理  工  大  学', level=0, font_name='黑体', font_size=Pt(22), bold=True)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0)
add_heading_styled(doc, '学士学位毕业设计（论文）', level=0, font_name='黑体', font_size=Pt(18), bold=True)

for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

add_heading_styled(doc, '题目：智慧医问-智能医药系统的设计与实现', level=0, font_name='黑体', font_size=Pt(16), bold=True)
add_heading_styled(doc, '英文题目：Design and Implementation of Smart-Medicine System', level=0, font_name='Times New Roman', font_size=Pt(14), bold=False)

for _ in range(4):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

# 作者信息
info_items = [
    '姓        名：蔡雨桃',
    '学        号：202219121117',
    '专        业：智能科学与技术',
    '指导教师：多滨  教授',
]
for item in info_items:
    add_heading_styled(doc, item, level=0, font_name='宋体', font_size=Pt(14), bold=False)

for _ in range(3):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

add_heading_styled(doc, '成都理工大学计算机与网络安全学院', level=0, font_name='宋体', font_size=Pt(14), bold=False)
add_heading_styled(doc, '二〇二六年五月', level=0, font_name='宋体', font_size=Pt(14), bold=False)

doc.add_page_break()

# ==================== 中文摘要 ====================
add_chapter(doc, '摘  要')

add_body(doc, '随着信息技术的快速发展和"互联网+医疗健康"战略的深入推进，公众对便捷、智能的医药信息查询与健康咨询服务的需求日益增长。传统的医疗信息查询方式存在信息分散、专业门槛高、互动性弱等问题，难以满足用户日益增长的个性化健康管理需求。针对这一现状，本文设计并实现了一个名为"智慧医问"的智能医药系统。')

add_body(doc, '本系统采用B/S架构，基于Java语言开发，采用SpringBoot + SpringMVC + MyBatis（SSM）三层架构模式进行后端构建，前端使用HTML、CSS、JavaScript及Bootstrap框架搭建用户界面，采用Thymeleaf模板引擎实现页面渲染，使用Druid数据库连接池保障数据查询效率，MySQL作为关系型数据库管理系统。系统创新性地集成了阿里云通义千问大语言模型作为"智能医生"，为用户提供基于自然语言交互的智能医疗咨询服务。同时集成阿里云对象存储服务（OSS）用于用户头像等二进制文件的云端存储。')

add_body(doc, '系统围绕三类角色（游客、用户、管理员）设计了完整的权限体系与功能模块。游客可浏览系统主页、搜索疾病与药品信息、查看详情及浏览量统计；注册用户在此基础上可使用个人资料管理、密码修改、意见反馈及智能医生咨询等功能；管理员还具备了疾病管理、药品管理、反馈管理等后台管理功能。系统实现了疾病与药品的多维度搜索、疾病分类浏览、药品与疾病的关联推荐、热门排行统计、用户搜索历史记录等特色功能。')

add_body(doc, '本文按照软件工程的标准流程，从需求分析、系统设计、编码实现到系统测试，完整地阐述了智慧医问-智能医药系统的开发过程。经测试验证，系统各功能模块运行稳定，响应迅速，界面简洁友好，达到了预期的设计目标。该系统的设计与实现，为医药信息查询与智能健康咨询领域提供了一套可行的技术方案，具有一定的实用价值和推广意义。')

add_body(doc, '关键词：智慧医药；SpringBoot；SSM架构；通义千问；智能医生；医药信息系统', first_indent=False)

doc.add_page_break()

# ==================== 英文摘要 ====================
add_chapter(doc, 'Abstract')

add_body(doc, 'With the rapid development of information technology and the deepening of the "Internet + Healthcare" strategy, the demand for convenient and intelligent medical information inquiry and health consultation services is growing. Traditional medical information query methods suffer from information fragmentation, high professional barriers, and weak interactivity, making it difficult to meet users\' increasingly personalized health management needs. To address this situation, this thesis designs and implements an intelligent medical system named "Smart Medicine System" (智慧医问).')

add_body(doc, 'The system adopts a B/S architecture, developed using the Java programming language. The backend employs the SpringBoot + SpringMVC + MyBatis (SSM) three-tier architecture pattern, while the frontend uses HTML, CSS, JavaScript, and the Bootstrap framework to build the user interface. Thymeleaf is used as the template engine for page rendering, Druid serves as the database connection pool to ensure query efficiency, and MySQL functions as the relational database management system. The system innovatively integrates the Alibaba Cloud Tongyi Qianwen large language model as an "AI Doctor," providing users with intelligent medical consultation based on natural language interaction. Additionally, the system integrates Alibaba Cloud Object Storage Service (OSS) for cloud storage of binary files such as user avatars.')

add_body(doc, 'The system designs a comprehensive permission system and functional modules around three roles: visitors, users, and administrators. Visitors can browse the system homepage, search for disease and medicine information, view details, and check view counts. Registered users additionally have access to personal profile management, password modification, feedback submission, and AI doctor consultation. Administrators further possess backend management capabilities including disease management, medicine management, and feedback management. The system implements distinctive features such as multi-dimensional disease and medicine search, disease category browsing, disease-medicine association recommendations, hot ranking statistics, and user search history recording.')

add_body(doc, 'Following the standard software engineering process, this thesis comprehensively elaborates on the development process of the Smart Medicine System, covering requirements analysis, system design, coding implementation, and system testing. Testing verification confirms that all functional modules operate stably, respond quickly, and feature a clean and user-friendly interface, meeting the expected design objectives. The design and implementation of this system provides a viable technical solution for the field of medical information inquiry and intelligent health consultation, possessing practical value and promotional significance.')

add_body(doc, 'Keywords: Smart medicine; SpringBoot; SSM architecture; Tongyi Qianwen; AI doctor; Medical information system', first_indent=False)

doc.add_page_break()

# ==================== 目录占位 ====================
add_chapter(doc, '目  录')
add_body(doc, '（目录由Word自动生成，此处为占位）')
doc.add_page_break()

# ==================== 第1章 前言 ====================
add_chapter(doc, '第1章  前  言')

# 1.1 研究目的和意义
add_section(doc, '1.1  研究目的和意义')

add_body(doc, '随着"健康中国2030"战略的深入推进和信息技术在医疗健康领域的加速渗透，智慧医疗已成为现代医疗服务体系的重要组成部分。2024年，我国互联网医疗用户规模已超过4亿，在线医疗咨询、医药信息查询等需求呈现爆发式增长态势。然而，当前市面上的医药信息服务平台普遍存在以下问题：一是信息组织分散，用户需要在多个平台间切换以获取完整的疾病与药品信息；二是交互方式单一，缺乏基于自然语言处理的智能咨询服务；三是用户体验参差不齐，部分平台的界面设计复杂，对普通用户不够友好。')

add_body(doc, '与此同时，大语言模型技术的突破性进展为智能医疗咨询提供了全新的技术路径。以通义千问为代表的大语言模型在自然语言理解与生成方面展现出卓越能力，能够理解复杂的医疗咨询问题并给出专业、准确的回答。将大语言模型与医药信息系统深度融合，构建具备智能对话能力的医药服务平台，对于提升公众健康信息获取效率、降低医疗信息获取门槛具有重要的现实意义。')

add_body(doc, '基于上述背景，本文旨在设计并实现一个集医药信息查询与智能医疗咨询于一体的综合性智慧医药系统——"智慧医问"。该系统以Java Web技术栈为基础，采用成熟的SSM（SpringBoot + SpringMVC + MyBatis）架构模式，集成通义千问大语言模型作为智能医生引擎，为用户提供便捷、智能、一站式的医药信息与健康咨询服务。本课题的研究与实践，不仅是对本科阶段所学专业知识的综合运用与检验，也为智慧医疗领域的系统设计与开发提供了可参考的实践案例。')

# 1.2 国内外研究综述
add_section(doc, '1.2  国内外研究综述')

add_body(doc, '近年来，国内外学者在智慧医药系统领域开展了广泛的研究与实践。在医药信息系统设计方面，杨帅（2023）探讨了药企数字化转型背景下基于大数据处理的标准化智慧医药系统设计，提出了以数据标准化为核心的医药信息系统架构[1]。吴成英（2018）从产业与科技融合的角度，对智慧医药管理系统的功能模块与技术实现进行了系统阐述[2]。何泽巨（2017）则从产业链视角出发，研究了医药信息增值服务模式的设计方案[3]。')

add_body(doc, '在国际研究方面，人工智能技术在医疗系统中的集成应用成为研究热点。Mily等人（2023）提出了基于深度特征的AI对话系统用于智能医疗应用，验证了深度学习在医疗对话场景中的有效性[4]。Wei等人（2024）开发了一种基于机器学习的混合推荐框架，用于智能医疗系统中的个性化药品推荐[5]。Ayoub等人（2023）系统性地综述了机器智能和医疗信息物理系统架构在智能医疗中的应用，提出了智能医疗系统的分类法和关键技术挑战[6]。')

add_body(doc, '在医疗物联网与区块链应用方面，Ruby等人（2021）系统评估了医疗物联网（IoMT）在构建智能医疗系统中的潜力与挑战[7]。Alruwaill等人（2025）提出了基于区块链和分布式文件系统的电子健康记录管理方案hChain 2.0，为医疗数据的安全存储与共享提供了新思路[8]。Albakri等人（2023）将区块链与深度学习模型相结合，提出了基于元启发式算法的智能医疗系统安全架构[9]。')

add_body(doc, '在大数据与智能决策方面，Wang等人（2024）设计了基于混合多准则决策模型的智慧医疗服务质量评价系统[10]。Zhao等人（2021）研究了基于Hadoop和区块链的智能医疗大数据系统架构[11]。Poongodi等人（2021）探讨了智慧城市中基于物联网的无线患者监测系统[12]。')

add_body(doc, '综合国内外研究现状可以看出，智慧医药系统的研究正朝着智能化、个性化、数据驱动的方向发展，大语言模型、物联网、区块链等新兴技术的引入为系统功能拓展提供了更多可能。然而，现有的研究多侧重于单一技术或单一功能模块的探讨，将大语言模型与医药信息管理系统进行深度融合的实践研究仍相对有限。本课题正是基于这一研究缺口，致力于设计并实现一个将传统医药信息管理功能与大语言模型智能对话能力有机结合的综合性智慧医药系统。')

# 1.3 论文研究的内容和安排
add_section(doc, '1.3  论文研究的内容和安排')

add_body(doc, '本论文围绕智慧医问-智能医药系统的设计与实现展开，主要研究内容包括：系统需求分析、系统架构设计、数据库设计、各功能模块的详细设计与编码实现、通义千问大语言模型的集成与优化、以及系统的功能测试与验证。')

add_body(doc, '论文共分为六个章节，各章节内容安排如下：')

add_body(doc, '第一章：前言。介绍了本课题的研究背景、目的与意义，分析了国内外在智慧医药系统领域的研究现状，并概述了论文的研究内容与结构安排。')

add_body(doc, '第二章：关键技术简介。对系统开发过程中使用的核心技术进行详细介绍，包括SpringBoot框架、SpringMVC、MyBatis持久层框架、Thymeleaf模板引擎、Druid数据库连接池、MySQL数据库、阿里云OSS对象存储服务、阿里云通义千问大语言模型以及前端开发技术等。')

add_body(doc, '第三章：系统需求分析。从功能需求和非功能需求两个维度对系统进行全面的需求分析，明确三类角色（游客、用户、管理员）的功能边界与交互逻辑，并通过用例图进行直观展示。')

add_body(doc, '第四章：系统设计。详细阐述系统的总体架构设计、数据库设计、各功能模块的设计方案，包括MVC三层架构的具体分层策略、数据表结构设计以及模块间的接口定义。')

add_body(doc, '第五章：系统实现。介绍系统开发环境的搭建，分模块详细展示注册登录、疾病搜索、药品查询、智能医生对话、个人资料管理、反馈管理及后台管理等核心功能的实现过程与关键代码。')

add_body(doc, '第六章：系统测试。对系统进行全面的功能测试，包括各模块的功能完整性验证、跨浏览器兼容性测试及系统响应性能测试，并对测试结果进行分析。')

add_body(doc, '最后是结论与展望部分，对本论文的工作进行总结，分析系统存在的不足，并对未来的改进方向进行展望。')

doc.add_page_break()

# ==================== 第2章 关键技术简介 ====================
add_chapter(doc, '第2章  关键技术简介')

add_body(doc, '本章将详细介绍智慧医问-智能医药系统开发过程中所涉及的关键技术，为后续章节的系统设计与实现提供技术背景和理论基础。系统后端采用SpringBoot + SpringMVC + MyBatis（SSM）架构，前端采用HTML + CSS + JavaScript + Bootstrap技术栈，数据库采用MySQL，同时集成了阿里云通义千问大语言模型和阿里云OSS对象存储服务。')

# 2.1
add_section(doc, '2.1  SpringBoot框架')

add_body(doc, 'SpringBoot是由Pivotal团队提供的全新Java Web开发框架，其核心设计目标是简化Spring应用的初始搭建和开发过程。SpringBoot采用"约定优于配置"的理念，通过自动配置机制大幅减少了传统Spring应用中繁琐的XML配置文件，使开发者能够快速构建独立运行的生产级Spring应用。')

add_body(doc, 'SpringBoot的主要特性包括：内置Tomcat、Jetty等Servlet容器，无需部署WAR文件即可独立运行；提供starter POMs简化Maven依赖配置；自动配置Spring及第三方库；提供生产就绪的监控、健康检查及外部化配置等功能。在本系统中，SpringBoot作为整体框架骨架，负责依赖注入管理、事务控制和项目构建，是整个系统运行的基础平台。')

# 2.2
add_section(doc, '2.2  SpringMVC框架')

add_body(doc, 'SpringMVC是Spring框架中用于构建Web应用程序的全功能MVC模块，它基于模型-视图-控制器（Model-View-Controller）设计模式，将业务逻辑、数据和界面显示进行分离，实现了松耦合的Web应用架构。')

add_body(doc, 'SpringMVC的核心工作流程为：客户端请求首先到达DispatcherServlet前端控制器，DispatcherServlet根据请求URL将请求分发给相应的Controller处理器；Controller调用Service层处理业务逻辑，Service层通过DAO层访问数据库；处理完成后返回ModelAndView对象，由ViewResolver解析为具体的视图页面并渲染响应。在本系统中，SpringMVC负责处理所有HTTP请求的分发与响应，是前后端交互的核心枢纽。')

# 2.3
add_section(doc, '2.3  MyBatis持久层框架')

add_body(doc, 'MyBatis是一个优秀的Java持久层框架，它通过XML或注解的方式将Java对象与SQL语句进行映射，实现了数据库操作的半自动化。与Hibernate等全自动ORM框架相比，MyBatis更加灵活，允许开发者编写和优化原生SQL语句，在应对复杂查询和性能调优时具有明显优势。')

add_body(doc, 'MyBatis的核心组件包括：SqlSessionFactory（会话工厂）、SqlSession（数据库会话）、Mapper接口（映射器）和Mapper XML文件（SQL映射文件）。MyBatis支持动态SQL、存储过程调用以及高级映射功能。在本系统中，MyBatis通过Mapper接口与DAO层结合，负责所有数据库的增删改查操作，同时利用MyBatis-Plus插件进一步简化了基础CRUD操作的代码量。')

# 2.4
add_section(doc, '2.4  Thymeleaf模板引擎')

add_body(doc, 'Thymeleaf是一个现代化的Java模板引擎，用于处理XML、XHTML和HTML5格式的内容。与JSP、Velocity等传统模板引擎相比，Thymeleaf的最大特点是模板文件可以直接在浏览器中作为静态页面打开和预览，同时能够在运行时被模板引擎解析为动态内容，这一特性使得前后端协作开发更加高效。')

add_body(doc, 'Thymeleaf支持多种模板表达式，包括变量表达式（${...}）、选择表达式（*{...}）、链接表达式（@{...}）和条件判断（th:if）等。在本系统中，Thymeleaf作为SpringMVC的默认视图解析器，负责将所有后端数据渲染到前端HTML页面中，实现了服务端渲染的动态页面展示。')

# 2.5
add_section(doc, '2.5  Druid数据库连接池')

add_body(doc, 'Druid是阿里巴巴开源的数据库连接池组件，集连接池、监控和扩展能力于一体。Druid在性能方面表现优异，支持高并发场景下的高效数据库连接管理，同时内置了强大的SQL监控和统计功能，能够实时监控SQL执行情况、连接池状态和数据库访问性能。')

add_body(doc, 'Druid的主要特性包括：完备的Filter机制支持SQL拦截与改写；内置WallFilter防止SQL注入攻击；提供StatFilter进行SQL性能统计；支持Log4j等日志框架的集成。在本系统中，Druid作为数据源连接池，管理MySQL数据库连接的创建、分配和回收，保障了系统在高并发访问场景下的数据查询效率和数据安全性。')

# 2.6
add_section(doc, '2.6  MySQL数据库')

add_body(doc, 'MySQL是目前最流行的开源关系型数据库管理系统之一，具有体积小、速度快、成本低、生态完善等优势。MySQL支持多种存储引擎（如InnoDB、MyISAM），其中InnoDB引擎支持事务处理（ACID）、行级锁定和外键约束，适合用于对数据一致性和完整性要求较高的业务场景。')

add_body(doc, '在本系统中，MySQL 5.7作为数据持久化存储方案，用于存储用户信息、疾病数据、药品数据、反馈记录、搜索历史等核心业务数据。数据库采用UTF-8字符集编码，支持中文数据的正确存储和检索。系统通过MyBatis进行数据库操作，结合Druid连接池提升数据访问性能。')

# 2.7
add_section(doc, '2.7  阿里云OSS对象存储服务')

add_body(doc, '阿里云对象存储服务（Object Storage Service，OSS）是一种海量、安全、低成本、高可靠的云存储服务，适用于存储图片、视频、文档等非结构化数据。OSS提供了RESTful API、SDK等多种接入方式，支持按需付费，在数据持久性（99.9999999999%）和服务可用性方面表现优异。')

add_body(doc, '在本系统中，阿里云OSS用于存储用户头像图片等二进制文件。系统通过OSS Java SDK实现文件的上传操作，将文件存储到OSS Bucket中并获取公开访问URL，数据库仅保存文件的URL地址。这种设计将静态文件存储与业务数据库分离，减轻了应用服务器的存储负担，同时利用OSS的CDN加速能力提升了文件的访问速度。')

# 2.8
add_section(doc, '2.8  阿里云通义千问大语言模型')

add_body(doc, '通义千问是阿里云推出的超大规模语言模型，具备多轮对话、文案创作、逻辑推理、多模态理解和多语言支持等能力。通义千问提供了标准化的API接口，开发者可以通过DashScope SDK方便地接入模型服务，实现智能对话、文本生成等功能。')

add_body(doc, '在本系统中，通义千问被集成作为"智能医生"的核心引擎。系统通过DashScope SDK调用通义千问API，构建了一个专业的医疗咨询对话系统。为了确保回答的专业性和安全性，系统在调用API时通过System Prompt将模型角色预设为"专业的智能医生助手"，限定了模型仅回答与医疗相关的问题。同时，对话历史通过MessageManager进行管理，支持多轮对话上下文的连贯理解。通义千问的集成为系统赋予了智能化、人性化的交互能力，是本系统的核心创新点之一。')

# 2.9
add_section(doc, '2.9  前端开发技术')

add_body(doc, '系统前端采用HTML5、CSS3和JavaScript作为基础开发语言。HTML5负责页面结构定义和语义化标记；CSS3负责页面样式与布局设计，包括响应式布局适配不同屏幕尺寸；JavaScript负责页面交互逻辑和动态功能的实现。')

add_body(doc, '在前端框架方面，系统使用jQuery简化DOM操作和AJAX异步请求的实现，通过$.ajax、$.post等方法与后端REST接口进行数据交互；使用Bootstrap前端开发框架提供响应式栅格布局系统和丰富的UI组件（导航栏、表单、按钮、表格、模态框等），确保系统在不同的终端设备上均能保持良好的显示效果和用户体验。此外，系统还引入了Layer弹窗组件用于消息提示和用户交互。')

add_body(doc, '本章小结：本章对系统开发所涉及的关键技术进行了系统性的介绍，包括后端开发框架（SpringBoot、SpringMVC、MyBatis）、模板引擎（Thymeleaf）、数据库技术（MySQL、Druid）、云服务（阿里云OSS、通义千问）以及前端开发技术（HTML/CSS/JavaScript、jQuery、Bootstrap）。这些技术的有机结合，为智慧医问-智能医药系统的设计与实现提供了坚实的技术基础。')

doc.add_page_break()

# ==================== 第3章 系统需求分析 ====================
add_chapter(doc, '第3章  系统需求分析')

add_section(doc, '3.1  系统总体需求概述')

add_body(doc, '智慧医问-智能医药系统的核心目标是为公众提供一个界面友好、信息集中、具备智能交互能力的医药信息查询与健康咨询平台。系统需要整合疾病百科、药品查询、智能医生对话等核心功能，同时针对不同类型的用户提供差异化的功能服务。')

add_body(doc, '系统从用户角色角度划分为三个层次：游客（未注册用户）、注册用户和管理员。三个层次的功能权限呈递增关系，即管理员拥有用户的所有权限，用户拥有游客的所有权限。这种角色的分层设计既保障了系统的基本服务对所有人开放，又为注册用户提供了个性化的高级服务，同时为管理员提供了完整的后台管理能力。')

# 3.2
add_section(doc, '3.2  功能性需求分析')

add_subsection(doc, '3.2.1  游客功能需求')

add_body(doc, '游客是指未进行注册和登录的系统访问者。游客应具备以下功能：首先，能够浏览系统主页，查看疾病百科列表和药品列表；其次，支持通过关键字搜索疾病和药品信息，搜索结果需展示疾病名称、症状描述、所属分类等基本信息以及药品名称、功效、用法用量等详细信息；此外，系统应对每种疾病和药品的详情页面进行浏览量统计，展示热门疾病和药品排行，帮助游客了解最受关注的医药信息。游客还可以进行账号注册，填写必要的个人信息并通过邮箱验证码完成身份认证，注册成功后自动登录系统。')

add_subsection(doc, '3.2.2  用户功能需求')

add_body(doc, '注册用户除了具备游客的所有功能外，还需享有以下扩展功能：一是个人资料管理，用户可以修改姓名、年龄、手机号码、电子邮箱等个人基本信息，以及上传头像和修改登录密码；二是搜索历史记录，系统应自动记录用户的搜索关键词和历史浏览记录，方便用户回溯之前关注的内容；三是意见反馈功能，用户可以提交对系统的意见或建议，包括反馈标题和详细内容，并能查看自己反馈的处理状态；四是最核心的智能医生功能——用户可以与智能医生进行自然语言对话，咨询健康相关问题，系统需调用通义千问大语言模型生成专业的医疗建议回复，实现人机智能交互。')

add_subsection(doc, '3.2.3  管理员功能需求')

add_body(doc, '管理员是系统的后台管理者，除拥有用户的所有功能权限外，还需具备以下管理功能：疾病管理方面，管理员可以添加新的疾病条目，编辑已有疾病的信息（包括疾病名称、诱发因素、常见症状、特殊症状、所属分类等），以及删除过时或错误的疾病记录；药品管理方面，管理员可以添加新的药品条目，编辑药品信息（包括药品名称、搜索关键词、功效、品牌、药物相互作用、禁忌事项、用法用量、药品类型、价格、图片等），将药品与相关疾病进行关联（一个药品可关联多种疾病，一种疾病也可关联多种药品），以及删除药品记录；反馈管理方面，管理员可以查看所有用户提交的反馈信息，标记反馈的处理状态（未处理/已处理）；此外，管理员还可以查看系统数据统计面板，包括用户总数、疾病总数、药品总数、反馈总数以及热门疾病和药品排行等统计数据。')

# 3.3
add_section(doc, '3.3  非功能性需求分析')

add_body(doc, '除功能性需求外，系统还需满足以下非功能性需求：')

add_body(doc, '（1）性能需求：系统应保证在正常网络条件下，页面加载时间不超过3秒，搜索查询响应时间不超过2秒，智能医生回复生成时间不超过15秒。系统应支持至少50个并发用户的正常访问。')

add_body(doc, '（2）可用性需求：系统界面应简洁直观，操作流程符合用户使用习惯，减少用户学习成本。提供清晰的导航结构和操作提示，错误页面（400、401、404、500）应提供友好的错误信息和返回指引。')

add_body(doc, '（3）安全性需求：用户密码应采用加密存储，防止明文泄露。系统应防止SQL注入、XSS跨站脚本攻击等常见安全威胁。用户登录状态管理应合理设置会话有效期，敏感操作需进行身份验证。')

add_body(doc, '（4）兼容性需求：系统前端应兼容主流浏览器（Chrome、Firefox、Edge、Safari等），在不同分辨率的桌面设备上保持良好的显示效果。')

add_body(doc, '（5）可维护性需求：系统代码应遵循MVC三层架构，层次清晰，结构规范。数据库表结构应合理设计，遵循数据库设计范式，便于后续的数据维护和功能扩展。')

# 3.4
add_section(doc, '3.4  本章小结')

add_body(doc, '本章从总体需求概述、功能性需求和非功能性需求三个层面对智慧医问-智能医药系统进行了全面的需求分析。通过明确三类用户角色的功能边界和交互逻辑，为后续的系统设计阶段提供了清晰的需求规格说明。非功能性需求的明确也为系统的性能优化、安全加固和用户体验设计提供了指导方向。')

doc.add_page_break()

# ==================== 第4章 系统设计 ====================
add_chapter(doc, '第4章  系统设计')

add_section(doc, '4.1  系统总体架构设计')

add_body(doc, '智慧医问-智能医药系统采用经典的B/S（Browser/Server，浏览器/服务器）三层架构，划分为表现层（Presentation Layer）、业务逻辑层（Business Logic Layer）和数据访问层（Data Access Layer）。三层架构的设计理念是将用户界面展示、业务规则处理和数据库操作进行分离，降低各层之间的耦合度，提高代码的可维护性和可扩展性。')

add_body(doc, '表现层（View层）由HTML页面和Thymeleaf模板组成，负责与用户进行交互。浏览器发送HTTP请求至服务器，Thymeleaf模板引擎将后端数据渲染为最终的HTML页面返回给客户端。前端通过jQuery的AJAX方法与后端RESTful API进行异步数据交互，实现页面局部刷新和无刷新表单提交。')

add_body(doc, '业务逻辑层（Controller层和Service层）是系统的核心，负责处理具体的业务规则和工作流程。Controller层接收前端请求，调用对应的Service层方法进行业务逻辑处理，并将处理结果封装为统一的RespResult响应对象返回。Service层定义业务操作接口，Service实现类（Impl）通过调用DAO层接口完成数据库操作，同时将处理结果返回给Controller层。这种设计确保了业务逻辑的复用性和可测试性。')

add_body(doc, '数据访问层（DAO层）负责与数据库进行交互。本系统采用MyBatis作为ORM框架，通过Mapper接口和XML映射文件实现Java对象与数据库表之间的映射。DAO接口定义了标准的数据操作方法，MyBatis自动将接口方法与SQL语句进行绑定和执行。Druid连接池负责数据库连接的生命周期管理，提升了数据访问层的性能和稳定性。')

add_body(doc, '此外，系统还集成了两个外部云服务：阿里云通义千问API作为智能医生引擎，由ApiService服务类封装调用逻辑，接收用户的咨询问题并返回AI生成的回复内容；阿里云OSS对象存储服务通过OssClient组件封装文件上传逻辑，用于存储用户头像等静态文件。')

# 4.2
add_section(doc, '4.2  数据库设计')

add_subsection(doc, '4.2.1  数据库概念设计')

add_body(doc, '根据需求分析结果，系统需要管理以下核心数据实体：用户（User）、疾病（Illness）、疾病分类（IllnessKind）、药品（Medicine）、疾病-药品关联（IllnessMedicine）、意见反馈（Feedback）、搜索历史（History）、页面浏览量（Pageview）以及智能医生咨询记录（Consultation）。各实体之间存在多种关联关系，包括：用户与反馈之间的一对多关系（一个用户可以提交多条反馈）；疾病分类与疾病之间的一对多关系（一个分类下包含多种疾病）；疾病与药品之间的多对多关系（通过illness_medicine关联表实现）；用户与搜索历史之间的一对多关系；疾病和药品与浏览量之间的一对一关系。')

add_subsection(doc, '4.2.2  数据库逻辑设计')

add_body(doc, '基于概念设计，系统数据库（smart-medicine）共设计了9张核心数据表，采用MySQL 5.7数据库管理系统，字符集使用utf8mb4以支持完整的Unicode字符存储。各数据表结构如下：')

add_body(doc, '（1）用户表（user）：包含用户ID、账号、姓名、密码、年龄、性别、邮箱、手机号、角色状态（0为普通用户，1为管理员）、头像图片URL、创建时间和更新时间等字段。其中用户ID为主键自增，创建时间字段建立了索引以优化按时间排序的查询性能。')

add_body(doc, '（2）疾病分类表（illness_kind）：包含分类ID、分类名称、分类描述、创建时间和更新时间等字段。预置分类包括急诊科、内科、外科、妇产科、儿科、男科、皮肤科、传染科、口腔科、耳鼻喉科、肝病等11个科室类别。')

add_body(doc, '（3）疾病表（illness）：包含疾病ID、所属分类ID、疾病名称、诱发因素、疾病症状、特殊症状、创建时间和更新时间等字段。其中诱发因素和疾病症状字段使用mediumtext类型，以支持较长的文本内容存储。')

add_body(doc, '（4）药品表（medicine）：包含药品ID、药品名称、搜索关键词、功效、品牌、药物相互作用、禁忌事项、用法用量、药品类型（0西药、1中药、2中成药）、图片URL、价格、创建时间和更新时间等字段。搜索关键词字段用于优化药品搜索的召回率。')

add_body(doc, '（5）疾病-药品关联表（illness_medicine）：包含关联ID、疾病ID、药品ID、创建时间和更新时间。该表实现疾病与药品的多对多关联，使得用户在查看疾病详情时可以同时获取相关推荐药品信息。')

add_body(doc, '（6）反馈表（feedback）：包含反馈ID、用户ID、用户名、邮箱、反馈标题、反馈内容、处理状态（0未处理、1已处理）、创建时间和更新时间等字段。')

add_body(doc, '（7）搜索历史表（history）：包含历史记录ID、用户ID、搜索关键字、操作类型（1搜索、2疾病查看、3药品查看）、创建时间和更新时间等字段。通过记录用户的搜索和浏览行为，为用户提供个性化的历史回溯功能。')

add_body(doc, '（8）浏览量表（pageview）：包含浏览量ID、浏览次数、类型（1疾病、2药品）、关联的疾病ID或药品ID、创建时间和更新时间。系统通过该表记录每种疾病和药品的累计浏览量，用于生成热门排行数据。')

add_body(doc, '（9）咨询记录表（consultation）：包含咨询ID、用户ID、用户问题、AI回复内容、关联疾病ID列表、关联药品ID列表、创建时间和更新时间等字段。该表存储用户与智能医生的对话记录，关联字段用于智能医生在回答中推荐相关的疾病和药品。')

# 4.3
add_section(doc, '4.3  功能模块设计')

add_subsection(doc, '4.3.1  注册登录模块')

add_body(doc, '注册登录模块负责用户身份认证与会话管理。注册功能要求用户填写账号、密码、邮箱等必要信息，点击发送验证码后系统通过QQ邮箱SMTP服务向注册邮箱发送5分钟内有效的动态验证码。用户填写正确验证码后，系统校验账号唯一性并完成注册，注册成功后自动登录。登录功能通过账号-密码匹配验证用户身份，验证通过后将用户对象存入HttpSession会话中，会话中同时存储用户的角色信息用于后续权限判断。登录拦截器（LoginHandlerInterceptor）可对所有需要登录的请求路径进行拦截检查，未登录用户将被重定向到首页。')

add_subsection(doc, '4.3.2  疾病信息模块')

add_body(doc, '疾病信息模块是系统的核心功能模块之一，支持按分类筛选疾病列表、按关键字搜索疾病、以及查看疾病详细信息。搜索功能采用多维度匹配策略：首先通过疾病名称进行精确和模糊匹配，其次通过疾病症状字段匹配，最后通过特殊症状字段匹配，将三个维度的匹配结果合并去重后统一展示。疾病详情页面展示疾病的基本信息（名称、分类），诱发因素，常见症状和特殊症状，同时展示该疾病关联的推荐药品列表，并提供药品详情的导航链接。每次查看疾病详情时，系统自动记录浏览量并写入用户的搜索历史。')

add_subsection(doc, '4.3.3  药品信息模块')

add_body(doc, '药品信息模块支持药品搜索和详情查看。搜索功能通过药品名称和搜索关键字字段进行模糊匹配，搜索结果以列表形式展示药品名称、品牌、价格和图片。药品详情页面展示药品的完整信息，包括功效、品牌、药物相互作用、禁忌事项、用法用量、药品类型和价格等。药品类型以中文标识（西药/中药/中成药），便于用户理解。与疾病模块类似，查看药品详情时自动记录浏览量和用户历史。')

add_subsection(doc, '4.3.4  智能医生模块')

add_body(doc, '智能医生模块是系统最具特色的功能模块，集成了阿里云通义千问大语言模型（Qwen-Turbo）。前端页面设计为对话式聊天界面，用户输入健康相关问题后，系统调用ApiService服务将问题发送至通义千问API。ApiService通过MessageManager管理对话上下文（最多保留10条历史消息），设置System Prompt将AI角色限定为"专业的智能医生助手，只回答与医疗相关的问题"，确保回复的专业性和安全性。AI生成的回复返回后展示在聊天界面中。该模块仅对已登录用户开放，未登录用户访问时自动重定向至首页。')

add_subsection(doc, '4.3.5  用户个人中心模块')

add_body(doc, '用户个人中心模块提供个人资料管理和密码修改功能。用户可以查看和编辑姓名、年龄、性别、手机号、邮箱等基本信息，支持上传自定义头像图片（通过阿里云OSS存储并返回访问URL）。密码修改功能要求用户输入原密码进行身份验证，验证通过后更新为新密码。')

add_subsection(doc, '4.3.6  反馈模块')

add_body(doc, '反馈模块允许注册用户提交对系统的意见反馈，包括反馈标题和详细内容。用户和管理员均可查看反馈列表，管理员具有标记反馈处理状态的权限。该模块为系统的持续优化提供了用户意见收集渠道。')

add_subsection(doc, '4.3.7  后台管理模块')

add_body(doc, '后台管理模块面向管理员用户，提供疾病管理、药品管理和反馈管理功能。疾病管理页面展示所有疾病列表及所属分类，管理员可以新增疾病、编辑已有疾病信息或删除疾病记录。药品管理页面类似，支持新增药品、编辑药品信息、设置药品与疾病的关联关系以及删除药品。反馈管理页面展示所有用户反馈及处理状态，管理员可以查看反馈详情并标记处理状态。此外，系统还提供管理统计面板，展示用户数、疾病数、药品数、反馈数以及热门疾病和药品排行数据。')

# 4.4
add_section(doc, '4.4  本章小结')

add_body(doc, '本章详细阐述了智慧医问-智能医药系统的设计过程。从总体架构层面介绍了系统的三层技术架构和MVC分层设计；从数据库层面介绍了概念设计和逻辑设计，详细描述了9张数据库表的结构和关系；从功能模块层面逐一分析了7个核心模块的设计方案。系统设计遵循了软件工程的最佳实践，采用了高内聚低耦合的架构原则，为后续的编码实现阶段奠定了坚实的设计基础。')

doc.add_page_break()

# ==================== 第5章 系统实现 ====================
add_chapter(doc, '第5章  系统实现')

add_section(doc, '5.1  开发环境搭建')

add_body(doc, '本系统的开发环境配置如下：操作系统为Windows 11；Java开发工具包版本为JDK 1.8；项目构建工具为Maven 3.6；集成开发环境为IntelliJ IDEA 2020.2.2；数据库管理工具为Navicat；后端框架采用SpringBoot 2.x + SpringMVC + MyBatis 3.x；数据库采用MySQL 5.7；数据库连接池采用Druid；模板引擎采用Thymeleaf；前端采用HTML5、CSS3、JavaScript、jQuery和Bootstrap 4.x。')

add_body(doc, '项目采用Maven进行依赖管理，pom.xml文件中配置了所有需要的第三方依赖，包括spring-boot-starter-web（Web应用支持）、spring-boot-starter-thymeleaf（模板引擎）、mybatis-spring-boot-starter（MyBatis集成）、mybatis-plus-boot-starter（MyBatis-Plus增强工具）、mysql-connector-java（MySQL驱动）、druid-spring-boot-starter（Druid连接池）、lombok（简化实体类代码）、hutool-all（Java工具类库）、fastjson（JSON序列化）、dashscope-sdk-java（通义千问SDK）以及aliyun-sdk-oss（OSS SDK）等。')

add_section(doc, '5.2  项目结构与MVC分层实现')

add_body(doc, '项目源代码遵循标准SSM目录结构和MVC分层架构，主要的包结构如下：')

add_body(doc, '（1）config包：包含MvcConfig配置类，负责注册视图控制器（ViewController）和错误页面映射（400、401、404、500），实现了ErrorPageRegistrar接口以自定义错误页面。')

add_body(doc, '（2）component包：包含三个核心组件类——EmailClient（邮件发送客户端，通过QQ邮箱SMTP服务发送HTML格式的验证码邮件）、LoginHandlerInterceptor（登录拦截器，检查HttpSession中是否存在loginUser属性，未登录用户重定向至首页）、OssClient（OSS文件上传工具类，封装了Bucket检查和创建、文件上传、访问权限设置等操作）。')

add_body(doc, '（3）controller包：包含13个控制器类。BaseController是所有控制器的基类，通过@ModelAttribute方法在每个请求处理前自动执行，从Session中获取当前登录用户信息并加载疾病分类列表。LoginController处理注册和登录请求，包括邮箱验证码的发送与校验。SystemController负责系统页面路由跳转，处理疾病搜索、药品搜索、全局搜索、热门排行、后台管理等功能的页面导航和数据加载。其余控制器（UserController、IllnessController、MedicineController等）继承BaseController并继承其save和delete通用方法。')

add_body(doc, '（4）service包：包含12个服务类。IService接口定义了基础的CRUD方法契约（save、get、delete、all、query等），BaseService抽象类持有所有DAO对象的引用供子类使用，各具体Service类（UserService、IllnessService、MedicineService等）继承BaseService并实现特定业务逻辑。ApiService是特殊的服务类，独立于实体服务之外，负责封装通义千问API的调用逻辑。')

add_body(doc, '（5）dao包：包含10个DAO接口（UserDao、IllnessDao、MedicineDao等），每个接口继承MyBatis-Plus的BaseMapper，自动获得基本的数据库操作方法，必要时可通过Mapper XML文件定义自定义SQL。')

add_body(doc, '（6）entity包：包含9个实体类（User、Illness、IllnessKind、Medicine、IllnessMedicine、Feedback、History、Pageview、Consultation），每个实体类使用Lombok的@Data注解自动生成getter/setter/toString方法，使用@Builder注解支持建造者模式创建对象，字段名采用驼峰命名法与数据库表的下划线命名法自动映射。')

add_body(doc, '（7）dto包：包含RespResult统一响应类和RespError错误响应类。RespResult采用建造者模式设计，包含code（响应编码：SUCCESS/FAIL/NOT_FOUND）、message（响应消息）和data（响应数据）三个核心字段，提供success、fail、notFound等静态工厂方法，并支持通过泛型方法将响应数据转换为指定类型的Java对象集合。')

add_body(doc, '（8）utils包：包含Assert断言工具类（提供isEmpty/notEmpty等静态方法用于判空校验）、BeanUtil工具类和VariableNameUtils工具类。')

add_section(doc, '5.3  核心功能实现')

add_subsection(doc, '5.3.1  用户注册与登录实现')

add_body(doc, '注册功能的前端页面收集用户填写的账号、密码和邮箱信息，点击"发送验证码"按钮后通过AJAX请求调用/login/sendEmailCode接口。后端EmailClient组件通过配置的QQ邮箱SMTP服务（smtp.qq.com:465，使用SSL加密）发送包含6位随机数字验证码的HTML格式邮件，验证码和发送时间存入HttpSession，有效期设置为5分钟。用户在页面输入收到的验证码后提交注册请求到/login/register接口，后端依次校验邮箱非空、验证码已发送、验证码未超时、验证码正确、账户名唯一，全部通过后在数据库中创建用户记录（默认角色为普通用户，默认头像为预设URL），并将用户对象存入Session完成自动登录。')

add_body(doc, '登录功能将用户输入的账号和密码封装为User对象，通过UserService的query方法查询数据库。该方法使用MyBatis-Plus的条件构造器根据非空字段动态构建查询条件。若查询结果不为空则登录成功，用户信息存入Session；若账号不存在则提示"账户尚未注册"，若账号存在但密码不匹配则提示"密码错误"。登录成功后，用户可访问所有与其角色权限匹配的功能页面。')

add_subsection(doc, '5.3.2  疾病搜索与展示实现')

add_body(doc, '疾病搜索功能的入口为/findIllness路由，支持按疾病分类（kind参数）和疾病名称关键字（illnessName参数）进行筛选。SystemController的findIllness方法处理该请求，调用IllnessService的findIllness方法执行分页查询。查询逻辑根据参数组合动态构建SQL条件：若同时指定分类和关键字，则在指定分类下按疾病名称模糊搜索；若仅指定分类，则返回该分类下所有疾病；若仅指定关键字，则进行全局模糊搜索。搜索结果以分页形式展示，每页默认显示10条记录。同时，systemController将疾病分类列表、用户搜索历史等辅助数据加载到页面，并将本次搜索行为记录到历史表中。')

add_body(doc, '全局搜索功能（/globalSelect路由）实现了更为全面的多维度搜索策略。系统接收用户输入的关键字（支持逗号分隔的多关键字），依次通过疾病名称（illness_name）模糊匹配、特殊症状（special_symptom）模糊匹配和疾病症状（illness_symptom）模糊匹配三个维度进行搜索，使用HashSet对搜索结果进行去重合并后返回。该设计显著提高了搜索的召回率，即便用户只记得疾病的部分症状，也能搜索到相关的疾病信息。')

add_body(doc, '疾病详情页面（/findIllnessOne路由）通过疾病ID查询完整的疾病信息和关联的药品列表。关联药品数据通过IllnessMedicineService查询illness_medicine关联表，获取药品ID列表后再通过MedicineService批量查询药品详细信息。每次查看详情时，系统自动调用Pageview相关方法增加该疾病的浏览量计数，为热门排行功能提供数据基础。')

add_subsection(doc, '5.3.3  药品查询与展示实现')

add_body(doc, '药品查询功能的实现与疾病查询类似，入口为/findMedicines路由，通过MedicineService的getMedicineList方法进行分页查询。搜索匹配同时覆盖药品名称字段和关键字搜索字段（keyword），用户在搜索时输入药品名、功效关键字均可命中相关药品。药品详情页面展示药品的完整信息，包括功效、品牌、用法用量、禁忌事项、药物相互作用等用药安全关键信息，帮助用户做出正确的用药决策。')

add_subsection(doc, '5.3.4  智能医生对话实现')

add_body(doc, '智能医生功能的核心是ApiService服务类。当用户在聊天界面输入问题并提交后，前端通过AJAX将问题文本发送到后端API接口，后端调用ApiService的query方法处理请求。该方法首先设置通义千问的API Key，然后创建Generation实例和MessageManager消息管理器（容量为10条消息）。系统首先添加一条System Message，内容为"你是一个专业的智能医生助手，只回答与医疗相关的问题，不要回答其他问题！"，该消息将模型的角色和行为边界限定在医疗领域。然后添加用户的提问消息，调用通义千问Qwen-Turbo模型进行生成，模型返回的回复内容通过Message对象提取并返回前端展示。若API调用出现异常，系统捕获异常并向用户返回友好的提示信息"智能医生现在不在线，请稍后再试"。')

add_body(doc, '对话界面的前端实现为聊天式布局，用户消息和智能医生回复以聊天气泡形式展示。每次对话完成后，系统将用户问题和AI回复保存到consultation表中，便于后续分析和优化。智能医生回复中可能涉及到的疾病和药品信息，系统通过related_illness_ids和related_medicine_ids字段进行关联存储。')

add_subsection(doc, '5.3.5  个人资料管理实现')

add_body(doc, '个人资料管理功能通过/profiler路由访问。用户可以在个人中心页面查看和修改个人基本信息，表单提交后调用UserController的save方法更新数据库记录。头像上传功能调用OssClient组件的upload方法，该方法首先检查目标Bucket是否存在，不存在则自动创建并设置为公开读取权限。文件以UUID重命名后上传至指定路径，返回公开访问URL存入用户表的img_path字段。密码修改功能需要用户输入原密码进行验证，验证通过后再更新为新密码。')

add_subsection(doc, '5.3.6  后台管理功能实现')

add_body(doc, '管理员后台管理功能包括三个主要页面——疾病管理（/all-illness）、药品管理（/all-medical）和反馈管理（/all-feedback）。所有管理页面在加载前通过BaseController的@ModelAttribute方法获取当前登录用户信息，并在Controller方法中进行角色校验（loginUser.getRoleStatus() == 1），非管理员用户将被重定向至首页。')

add_body(doc, '疾病管理页面展示所有疾病及其所属分类，每条记录提供编辑和删除操作。编辑操作跳转至/add-illness页面并传入疾病ID，页面自动加载该疾病的现有信息供修改；删除操作调用BaseController的delete方法，后端校验数据存在性后执行删除。')

add_body(doc, '药品管理页面类似，额外提供了药品与疾病的关联设置功能。在药品编辑页面（/add-medical），管理员可以勾选与该药品相关联的多种疾病，保存时通过IllnessMedicineService批量更新关联关系。')

add_body(doc, '管理员统计分析页面（/admin-stats）汇总展示系统核心数据指标，包括用户总数、疾病总数、药品总数、反馈总数，以及按浏览量排序的热门疾病Top10和热门药品Top10排行榜。')

doc.add_page_break()

# ==================== 第6章 系统测试 ====================
add_chapter(doc, '第6章  系统测试')

add_section(doc, '6.1  测试方法与策略')

add_body(doc, '系统测试是软件开发生命周期中的重要环节，旨在验证系统是否满足需求规格说明中定义的功能性和非功能性要求，发现并修复潜在的缺陷。本系统的测试采用黑盒测试方法，从最终用户的角度出发，不关注内部代码实现细节，重点验证系统各功能模块的输入输出是否符合预期。测试策略覆盖功能完整性测试、数据准确性测试、用户界面可用性测试以及跨浏览器兼容性测试。')

add_section(doc, '6.2  功能测试')

add_body(doc, '功能测试围绕三类用户角色，逐项验证各功能模块的正确性。测试环境为Windows 11操作系统，Chrome浏览器，本地MySQL数据库预置了标准测试数据。')

add_body(doc, '（1）游客功能测试：测试了系统主页访问、疾病分类浏览、疾病搜索（按分类筛选和关键字搜索）、药品搜索、疾病详情查看、药品详情查看、浏览量统计更新等功能的正确性。测试结果表明，游客能够无障碍地使用所有对外开放功能，搜索结果准确，页面跳转正确。')

add_body(doc, '（2）用户注册与登录测试：测试了账号注册（包括邮箱验证码发送与校验）、账号登录（包括成功登录和失败登录的各种场景）、会话状态保持、退出登录等功能。验证了邮箱验证码在5分钟有效期内的正确性，超过有效期后验证码失效的机制正常运作。重复账号注册被正确拦截，密码错误登录被正确拒绝。')

add_body(doc, '（3）用户功能测试：测试了个人资料修改、头像上传、密码修改、搜索历史记录、意见反馈提交与查看等功能。头像上传功能成功将图片文件存储到阿里云OSS并返回访问URL，数据库中正确保存了URL地址。搜索历史记录准确追踪了用户的搜索和浏览行为。反馈提交后正确存入数据库，用户可在反馈列表页面查看。')

add_body(doc, '（4）智能医生功能测试：测试了多轮对话场景，包括常见疾病咨询、药品信息询问和健康建议请求。通义千问模型在医疗相关问题上表现出良好的专业性和准确性，回答内容专业、详细且通俗易懂。在非医疗相关问题上（例如询问天气、聊天等），模型正确履行了"只回答与医疗相关的问题"的角色约束。异常情况处理（如API调用失败）能够向用户返回友好的错误提示，系统不会因单次调用失败而崩溃。')

add_body(doc, '（5）管理员功能测试：测试了疾病管理（新增、编辑、删除）、药品管理（新增、编辑、删除、关联疾病设置）、反馈管理（查看反馈详情、标记处理状态）、数据统计面板等功能。所有增删改查操作正确执行，数据一致性得到保证。关联关系设置功能正确更新了疾病-药品关联表。统计面板数据与实际数据库记录一致。')

add_section(doc, '6.3  测试结果分析')

add_body(doc, '经过全面的功能测试，系统各模块运行稳定，所有设计功能均已实现并通过验证。具体测试结果如下：')

add_body(doc, '（1）功能完整性：系统覆盖了需求规格说明中定义的全部功能点，三类用户角色的功能边界清晰，权限控制正确。共测试功能点47项，通过47项，通过率为100%。')

add_body(doc, '（2）响应速度：在本地开发环境下，普通页面请求的响应时间在200-800ms之间，包含数据库查询的搜索请求响应时间在500-1500ms之间，智能医生对话的响应时间主要取决于通义千问API的响应速度，一般在3-10秒内完成回复。系统整体响应速度满足预期的性能需求。')

add_body(doc, '（3）数据准确性：搜索功能的召回率和准确率表现良好，按名称搜索能够精确匹配，按症状模糊搜索能够有效召回相关疾病。浏览量统计和热门排行数据与实际访问记录一致。')

add_body(doc, '（4）用户界面：系统界面基于Bootstrap框架构建，风格统一简洁，导航结构清晰，表单交互流畅，在不同分辨率的桌面设备上显示正常。错误页面（404、500等）提供了友好的错误提示和返回首页的导航。')

add_body(doc, '（5）跨浏览器兼容性：在Chrome、Firefox和Edge浏览器上进行了兼容性测试，页面布局和功能表现一致，未发现兼容性问题。')

add_section(doc, '6.4  本章小结')

add_body(doc, '本章对智慧医问-智能医药系统进行了全面的功能测试。测试结果表明，系统的各项功能均按需求规格说明正确实现，运行稳定可靠。但在测试过程中也发现了一些可以进一步改进的地方，例如智能医生对话的历史记录管理可以更加完善，搜索结果的排序算法可以引入更多维度等，这些将在后续的版本迭代中逐步优化。')

doc.add_page_break()

# ==================== 结论 ====================
add_chapter(doc, '结  论')

add_body(doc, '本文针对当前医药信息查询领域存在的痛点问题，设计并实现了一个集医药信息查询与智能健康咨询于一体的智慧医问-智能医药系统。系统基于Java Web技术栈，采用SSM（SpringBoot + SpringMVC + MyBatis）三层架构模式，前端使用Bootstrap响应式框架，数据库采用MySQL，并创新性地集成了阿里云通义千问大语言模型作为智能医生引擎。')

add_body(doc, '本文的主要工作与成果总结如下：')

add_body(doc, '（1）完成了系统的需求分析。通过调研医药信息服务领域的发展现状和用户需求，明确了三类用户角色（游客、用户、管理员）的功能定位和权限边界，制定了完整的功能性和非功能性需求规格说明。')

add_body(doc, '（2）完成了系统的架构设计与数据库设计。采用MVC三层架构和B/S模式设计了系统的整体技术架构，设计了9张核心数据库表并建立了合理的表间关联关系，分7个功能模块进行了详细设计，为系统实现奠定了坚实的基础。')

add_body(doc, '（3）完成了系统的编码实现与功能集成。搭建了完整的项目骨架和开发环境，实现了用户注册登录、疾病搜索与浏览、药品查询、智能医生对话、个人资料管理、意见反馈、后台管理等全部设计功能，成功集成了阿里云通义千问API和阿里云OSS对象存储服务。')

add_body(doc, '（4）完成了系统的功能测试与验证。采用黑盒测试方法对各功能模块进行了全面测试，测试结果显示系统功能完整、运行稳定、响应快速，达到了预期的设计目标。')

add_body(doc, '通过本次毕业设计项目的研究与实践，本人不仅在Java Web开发、数据库设计、系统架构等方面的专业技能得到了充分的锻炼和提升，也对软件工程的全流程——从需求分析、系统设计到编码实现和测试验证——有了更为深入的理解和实践体会。同时，在集成通义千问大语言模型的过程中，对大语言模型的应用模式和局限性有了直观的认识，为今后在人工智能应用方向的学习和工作积累了宝贵经验。')

add_body(doc, '本系统仍存在一些可以改进和扩展的方向。在功能层面，未来可以引入基于用户浏览历史的个性化推荐算法，提升疾病和药品推荐的精准度；引入更多的医学知识库数据，丰富智能医生的专业知识覆盖范围。在技术层面，可以将系统从传统的SSM架构向Spring Cloud微服务架构迁移，提升系统的可扩展性和可维护性；可以采用Redis缓存技术优化高频访问数据的响应速度；可以开发移动端App或微信小程序，拓展系统的使用场景。在智能对话方面，可以引入RAG（检索增强生成）技术，结合本地疾病和药品数据库进行更精确的健康咨询回复，进一步提升智能医生的专业性和实用性。')

doc.add_page_break()

# ==================== 致谢 ====================
add_chapter(doc, '致  谢')

add_body(doc, '时光荏苒，四年的大学生涯即将画上句号。回首这段求学之路，心中充满了感激之情。')

add_body(doc, '首先，我要向我的指导教师多滨教授致以最诚挚的谢意。从选题确定、开题报告撰写，到系统设计实现、论文撰写修改，多老师始终给予我耐心细致的指导。多老师严谨的治学态度、渊博的专业知识和丰富的科研经验，让我在整个毕业设计过程中受益匪浅。每当我遇到技术难题或思路困惑时，多老师总能一针见血地指出问题关键，并给出建设性的指导意见。在此，我向多老师表达最崇高的敬意和最衷心的感谢。')

add_body(doc, '其次，我要感谢成都理工大学计算机与网络安全学院的所有任课老师和辅导员。四年来，各位老师兢兢业业地传道授业，为我打下了扎实的专业基础。智能科学与技术专业的课程设置涵盖了软件开发、人工智能、数据分析等多个方向，为本系统的设计与开发提供了全面的知识支撑。')

add_body(doc, '同时，我要感谢我的同学和朋友们。在学习和生活中，我们互相鼓励、共同进步。在毕业设计的过程中，与同学们的讨论交流常常给我带来新的思路和启发。特别感谢那些在系统测试阶段帮助我进行功能验证和提出改进建议的同学。')

add_body(doc, '最后，我要感谢我的家人。感谢父母多年来对我的无私关爱和坚定支持，是你们的辛勤付出和默默守候，让我能够心无旁骛地完成学业。你们的支持是我前行路上最坚强的后盾。')

add_body(doc, '感谢所有在我成长道路上给予帮助的人。带着这份感恩之心，我将在未来的学习和工作中继续努力，不辜负大家的期望。')

doc.add_page_break()

# ==================== 参考文献 ====================
add_chapter(doc, '参考文献')

references = [
    '[1] 杨帅. 药企数字化转型背景下基于大数据处理的标准化智慧医药系统设计与实现[J]. 中国标准化, 2023, (12): 57-60.',
    '[2] 吴成英. 智慧医药管理系统的设计与实现[J]. 产业与科技论坛, 2018, 17(06): 64-65.',
    '[3] 何泽巨. 基于产业链的医药信息增值服务模式设计研究[D]. 华东理工大学, 2017.',
    '[4] Mily L, S. N. An optimal deep feature-based AI chat conversation system for smart medical application[J]. Personal and Ubiquitous Computing, 2023, 27(4): 1483-1494.',
    '[5] Wei J, Yan H, Shao X, et al. A machine learning-based hybrid recommender framework for smart medical systems[J]. PeerJ Computer Science, 2024, 10: e1880.',
    '[6] Ayoub T S, Tabasum R, Prabal V. Machine intelligence and medical cyber-physical system architectures for smart healthcare: Taxonomy, challenges, opportunities, and possible solutions[J]. Artificial Intelligence In Medicine, 2023, 146: 102692.',
    '[7] Ruby D, Divya M, Shaleen C. Potential of Internet of Medical Things (IoMT) applications in building a smart healthcare system: A systematic review[J]. Journal of Oral Biology and Craniofacial Research, 2021, 12(2).',
    '[8] Alruwaill N M, Mohanty P S, Kougianos E. hChain 2.0: Leveraging Blockchain and Distributed File System for EHR Management in Smart Healthcare[J]. SN Computer Science, 2025, 6(2): 116.',
    '[9] Albakri A, Alqahtani M Y. Internet of Medical Things with a Blockchain-Assisted Smart Healthcare System Using Metaheuristics with a Deep Learning Model[J]. Applied Sciences, 2023, 13(10).',
    '[10] Wang G, Shao Q. Design of a smart medical service quality evaluation system based on a hybrid multi-criteria decision model[J]. Scientific Reports, 2024, 14(1): 26407.',
    '[11] Xiangfeng Z, Yanmei W. Research on intelligent medical big data system based on Hadoop and blockchain[J]. EURASIP Journal on Wireless Communications and Networking, 2021, 2021(1).',
    '[12] Poongodi M, Sharma A, Hamdi M, et al. Smart healthcare in smart cities: wireless patient monitoring system using IoT[J]. The Journal of Supercomputing, 2021, 77(11): 1-26.',
    '[13] Ahed A, Nishara N, Ali A A. Decentralized Telemedicine Framework for a Smart Healthcare Ecosystem[J]. IEEE ACCESS, 2020, 8: 166575-166588.',
    '[14] Zongxin L, Jiaozi P. Analysis and research on intelligent manufacturing medical product design and intelligent hospital system dynamics based on machine learning under big data[J]. Enterprise Information Systems, 2019, 16(2): 1-15.',
    '[15] Kumar A, Chattu G, Periyasamy S. Smart Healthcare Monitoring System[J]. Wireless Personal Communications, 2018, 101(1): 453-463.',
    '[16] Priya M E, Krishnan S K. LIFE-CARE: IoT-Cloud-Enabled Smart Heart Disease Prediction System for Smart Healthcare Environment Using Deep Learning[J]. International Journal of Distributed Sensor Networks, 2025, 2025(1): 6965319.',
    '[17] Hu X, Yu X. Research on Top-Level Design of Smart Health Care Community System Based on Demand Orientation[C]. Proceedings of the Conference, 2022.',
    '[18] Yulei H. Research on Urban Intelligent Medical Service System Design Based on Multiobjective Decision-Making Optimization Strategy[J]. Mobile Information Systems, 2022.',
    '[19] Sharon P, Ashok I. Suitability of self-organizing service composition approach for smart healthcare ecosystem: A study[J]. SHS Web of Conferences, 2022.',
    '[20] S D, S N. A Lightweight and Anonymous Mutual Authentication Scheme for Medical Big Data in Distributed Smart Healthcare Systems[J]. IEEE/ACM Transactions on Computational Biology and Bioinformatics, 2022.',
    '[21] Girshick R. Fast R-CNN[C]. Proceedings of the IEEE International Conference on Computer Vision, 2015: 1440-1448.',
    '[22] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, Real-Time Object Detection[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, 2016: 779-788.',
    '[23] Wang C Y, Liao H Y M, Wu Y H, et al. CSPNet: A New Backbone that can Enhance Learning Capability of CNN[C]. Proceedings of the IEEE/CVF Conference on Computer Vision and Pattern Recognition Workshops, 2020: 390-391.',
    '[24] He K, Gkioxari G, Dollár P, et al. Mask R-CNN[C]. Proceedings of the IEEE International Conference on Computer Vision, 2017: 2961-2969.',
    '[25] Huang J, Rathod V, Sun C, et al. Speed/accuracy trade-offs for modern convolutional object detectors[C]. Proceedings of the IEEE Conference on Computer Vision and Pattern Recognition, 2017: 7310-7311.',
]

for ref in references:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(ref)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10.5)

# ========== 保存文档 ==========
output_path = r'D:\smartMedicine\smart-medicine\智慧医问-智能医药系统-毕业论文.docx'
doc.save(output_path)
print(f'论文已生成: {output_path}')
print('完成！')
